import sys
import subprocess
import os

# Install python-docx if not installed
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

def create_docx_from_md(md_file, docx_file):
    document = Document()
    
    # Add a title style modification if needed
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if not line and not in_code_block:
            continue

        if in_code_block:
            p = document.add_paragraph(line)
            p.style = 'No Spacing'
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            continue

        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            p = document.add_paragraph(line[2:], style='List Bullet')
        else:
            # Handle bold text **...** roughly
            p = document.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: # Odd parts are inside ** **
                    run.bold = True

    document.save(docx_file)
    print(f"Successfully created {docx_file}")

if __name__ == "__main__":
    create_docx_from_md("PROJECT_DOCUMENTATION.md", "Project_One_Documentation.docx")
